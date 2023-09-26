from pingouin import cronbach_alpha
from factor_analyzer import FactorAnalyzer
from docx import Document
import pandas as pd
import numpy as np
import seaborn as sns
import statsmodels.api as sm
from sklearn.preprocessing import LabelEncoder
from semopy import Model as SEM_Model


import os


output_path = './output/'
# Check if the directory exists
if not os.path.exists(output_path):
    # If it doesn't exist, create the directory
    os.makedirs(output_path)
    print(f"Directory '{output_path}' created.")

def read_excel_to_dict(excel_path):
    # Read the Excel workbook
    xls = pd.ExcelFile(excel_path)
    
    # Initialize an empty dictionary to store all the data
    all_data = {}
    
    # Loop through each sheet and read data into a DataFrame, then convert it to a dictionary
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name)
        all_data[sheet_name] = df.to_dict(orient='index')
        
    return all_data

# Display column names with index
def display_col_name_index(df):
  # Iterate through columns and display column names with indexes in brackets
  for i, column_name in enumerate(df.columns):
      print(f"Column {i} [{column_name}]")

# Rename columns for demographic data columns: to have shorter column names
def rename_demographic_columns_by_index(df, demographic_cols):
    demographic_cols_names = []  # Initialize a list to store independent column names
    for col in demographic_cols:
        col_index = col.get('col_index')
        new_column_name = col.get('name')
        df.rename(columns={df.columns[col_index]: new_column_name}, inplace=True)
        demographic_cols_names.append(new_column_name)
    return df, demographic_cols_names

# Rename Independent columns according to Independent_Cols
def rename_independent_columns(df, independent_cols):
    independent_cols_names = []  # Initialize a list to store independent column names
    for col in independent_cols:
        col_name = col.get('name')
        col_index_from = col.get('col_index_from')
        col_index_to = col_index_from + col.get('number_questions')

        for i in range(col_index_from, col_index_to):
            old_col_name = df.columns[i]
            new_col_name = f'{col_name}_Q{i - col_index_from + 1}'
            df.rename(columns={old_col_name: new_col_name}, inplace=True)
            independent_cols_names.append(new_col_name)  # Append the new column name

    return df, independent_cols_names

# Rename Dependent/Target columns according to Dependent_Cols
def rename_target_columns(df, target_cols):
    target_cols_names = []  # Initialize a list to store Target column names
    for col in target_cols:
        col_name = col.get('name')
        col_index_from = col.get('col_index_from')
        col_index_to = col_index_from + col.get('number_questions')
        # Rename the column
        for i in range(col_index_from, col_index_to):
            old_col_name = df.columns[i]
            new_col_name = f'{col_name}_Q{i - col_index_from + 1}'
            df.rename(columns={old_col_name: new_col_name}, inplace=True)
            target_cols_names.append(new_col_name)  # Append the new column name

    return df, target_cols_names

# Create a statistics table for selected columns
def compute_selected_cols_statistics(df, selected_cols_names, output_file_name):
    nbr_data_points_cleaned = df.shape[0]
    stats_table = {}
    for col_name in selected_cols_names:
        freq = df[col_name].value_counts().sort_index()
        percent = (freq / nbr_data_points_cleaned) * 100
        percent = percent.round(1)  # Round to two decimal places
        stats_table[col_name] = pd.DataFrame({'Frequency': freq, 'Percent': percent})
        
        # Sort the DataFrame by 'Frequency' in descending order
        stats_table[col_name] = stats_table[col_name].sort_values(by='Frequency', ascending=False)
    
    # Save to Excel file
    output_file_name = f'{output_file_name}_Stats'
    save_to_excel(output_file_name, stats_table)

    del df
    return stats_table


# Convert the data in a nested dict into a dictionary of DataFrames
def convert_to_dataframe(table):
    dfs = {}
    for key, value in table.items():
        df = value
        dfs[key] = df

    # Concatenate the DataFrames vertically (along rows)
    combined_df = pd.concat(dfs.values(), axis=1, ignore_index=True)

    return combined_df

# Save the statistics table to an Excel file
def save_to_excel(filename, table):
    with pd.ExcelWriter(f'{output_path}{filename}.xlsx') as writer:
        for col, table in table.items():
            table.to_excel(writer, sheet_name=f'{col}')

# Extract data for selected columns
def extract_selected_colums_data(df, selected_cols_names):
    data_to_extract = df[selected_cols_names]
    data_to_extract = convert_likert_to_numerical(data_to_extract, extended_likert_mapping_all_languages)
    return data_to_extract

# Convert values of dataframe to numeric
def convert_values_to_numeric(df):
    # Initialize the LabelEncoder
    label_encoder = LabelEncoder()

    # Encode each column in the DataFrame
    for column in df.columns:
        df[column] = label_encoder.fit_transform(df[column])
    return df

# Function to manually convert Likert-scale responses to numerical values
# Function to manually convert Likert-scale responses to numerical values
def convert_likert_to_numerical(df, likert_mapping):
    df_numerical = df.copy()
    # Convert all columns to string type for consistent mapping
    df_numerical = df_numerical.astype(str)
    # Convert the mapping keys to string type as well
    str_likert_mapping = {str(key): value for key, value in likert_mapping.items()}
    
    for column in df.columns:
        df_numerical[column] = df_numerical[column].str.lower().map(str_likert_mapping)
        
    # Check for any remaining non-numeric values
    if df_numerical.applymap(np.isreal).all().all() == False:
        raise AssertionError("Not all columns are numeric after mapping. Check your mapping and DataFrame.")
        
    del df
    return df_numerical




# Sample Likert scale mapping (you can adjust this based on your actual survey Likert scale)
extended_likert_mapping_all_languages = {
    # English Labels
    'strongly disagree': 1,
    'disagree': 2,
    'neutral': 4,
    'agree': 6,
    'strongly agree': 7,
    'no opinion': 4,
   
    # Vietnamese Labels
    'hoàn toàn không đồng ý': 1,
    'không đồng ý': 2,
    'không ý kiến': 4,
    'đồng ý': 6,
    'hoàn toàn đồng ý': 7,
    'trung lap': 4,
    
    # Numerical Levels
    1: 1,
    2: 2,
    3: 3,
    4: 4,
    5: 5,
    6: 6,
    7: 7,
    8: 8,
    9: 9,
}

# Convert dataframe to numerical, keeping ordinality
def convert_df_to_numerical(df):
    # Initialize a list to store the names of detected ordinal columns
    detected_ordinal_columns = []

    # Check each column in the survey data
    for column in df.columns:
        # Get the unique values in the column (excluding NaN)
        unique_values = set(df[column].dropna().unique())
        
        # Check if the unique values are a subset of the extended ordinal mapping keys
        if unique_values.issubset(set(extended_likert_mapping_all_languages.keys())):
            # Column is detected as ordinal, mark it as an ordinal column
            detected_ordinal_columns.append(column)
                
            # Apply the extended ordinal mapping to convert the ordinal responses to numerical values
            df[column] = df[column].map(extended_likert_mapping_all_languages)

    # Identify categorical columns (excluding the detected ordinal columns)
    categorical_columns = df.select_dtypes(include=['object']).columns.difference(detected_ordinal_columns)

    # Convert categorical columns to numerical values using one-hot encoding
    df = pd.get_dummies(df, columns=categorical_columns)

    return df



# Remove data points (rows) with missing data related to selected columns
def remove_rows_with_missing_data_related_to_selected_cols(df, selected_cols_names):
    total_num_data_points = df.shape[0]
    # Filter rows with missing data only in the selected columns
    df = df.dropna(subset=selected_cols_names)
    # Number of data points removed
    number_datapoints_removed = total_num_data_points - df.shape[0]
    print(f'\nNumber of data points removed: \n{number_datapoints_removed}')

    return df

# Testing the reliability of the scale using Cronbach's alpha for independent columns
def compute_cronbach_alpha(selected_data, independent_cols):
    # Calculate Cronbach's alpha
    overall_alpha = cronbach_alpha(selected_data)
    overall_alpha_evaluation = interpret_alpha(overall_alpha[0])

    # Compute Cronbach Alpha for each Independent Variables
    alpha_table = []
    for col in independent_cols:
        col_names = []
        variable_name = col.get('name')
        nbr_questions = col.get('number_questions')
        for idx in range(nbr_questions):
            col_names.append(f'{variable_name}_Q{idx+1}')
        # print(col_names)
        data_of_this_variable = selected_data[col_names]
        # Calcul Alphe for this variable
        alpha_variable = cronbach_alpha(data_of_this_variable)
        alpha_evaluation = interpret_alpha(alpha_variable[0])
        alpha_table.append([alpha_variable, alpha_evaluation])

    # Print the result
    print(f"Overall Cronbach's Alpha: {overall_alpha}")

    return f"{overall_alpha}, '{overall_alpha_evaluation}'", alpha_table

def interpret_alpha(alpha):
    if alpha > 0.9:
        return "Excellent"
    elif alpha > 0.8:
        return "Good"
    elif alpha > 0.7:
        return "Acceptable"
    elif alpha > 0.6:
        return "Questionable"
    elif alpha > 0.5:
        return "Poor"
    else:
        return "Unacceptable"

# EFA analysis
def do_efa_analysis(selected_data, num_factors):
    fa = FactorAnalyzer(n_factors=num_factors, rotation='varimax')
    fa.fit(selected_data)
    factor_loadings = fa.loadings_
    # print(f"EFA loadings:\n{loadings}")
    # Visualize the factor loadings
    loadings_df = pd.DataFrame(factor_loadings, columns=[f'Factor {i+1}' for i in range(num_factors)], index=selected_data.columns)
    # Save the EFA analysis to an Excel file
    filename = 'EFA_Analysis'
    # Set index=False to exclude the DataFrame index
    filename = 'EFA_Analysis'
    excel_file_path = f'{output_path}{filename}.xlsx'
    loadings_df.to_excel(excel_file_path, index=True)  

    return loadings_df

def interpret_based_on_loadings(loadings_df, threshold=0.5):
    interpretations = []
    for factor in loadings_df.columns:
        strong_vars = loadings_df.loc[loadings_df[factor].abs() >= threshold].index.tolist()
        
        if strong_vars:
            sentence = f"In simple terms, '{factor}' is mainly about {', '.join(strong_vars)}."
            interpretations.append({'Factor': factor, 'Interpretation': sentence})

    interpretation_df = pd.DataFrame(interpretations)
    return interpretation_df

def recommendations_based_on_loadings(loadings_df, threshold=0.5):
    recommendations = []
    for factor in loadings_df.columns:
        strong_vars = loadings_df.loc[loadings_df[factor].abs() >= threshold].index.tolist()
        
        if strong_vars:
            advice = f"To improve in '{factor}', consider focusing on {', '.join(strong_vars)}."
            recommendations.append({'Factor': factor, 'Recommendation': advice})

    recommendation_df = pd.DataFrame(recommendations)
    return recommendation_df

    
# Pearson correlation matrix
def compute_correlation(selected_data):
    correlation_table = selected_data.corr(method='pearson')
    print("Pearson correlation coefficient matrix: ")
    print(correlation_table)
    # Save the Correlation Table to an Excel file
    # Set index=False to exclude the DataFrame index
    filename = 'Correlation_Table'
    excel_file_path = f'{output_path}{filename}.xlsx'
    correlation_table.to_excel(excel_file_path, index=True)  

    return correlation_table

def interpret_and_recommend_correlation(correlation_table, threshold=0.5):
    interpretations = []
    recommendations = []

    # Loop over the upper triangle of the correlation matrix
    for i, row_var in enumerate(correlation_table.index):
        for j, col_var in enumerate(correlation_table.columns[i+1:]):  # i+1 to skip self-correlation
            corr_value = correlation_table.loc[row_var, col_var]

            # Only consider strong correlations (positive or negative)
            if abs(corr_value) >= threshold:
                if corr_value > 0:
                    interpretation = f"{row_var} and {col_var} have a strong positive relationship."
                    recommendation = f"Improving {row_var} could also positively impact {col_var}, and vice versa."
                else:
                    interpretation = f"{row_var} and {col_var} have a strong negative relationship."
                    recommendation = f"Focusing on {row_var} might have the opposite effect on {col_var}, so be cautious."

                interpretations.append({
                    'Variables': f"{row_var} & {col_var}",
                    'Correlation': corr_value,
                    'Interpretation': interpretation
                })

                recommendations.append({
                    'Variables': f"{row_var} & {col_var}",
                    'Recommendation': recommendation
                })

    interpretation_df = pd.DataFrame(interpretations)
    recommendation_df = pd.DataFrame(recommendations)

    return interpretation_df, recommendation_df

# Regression Analysis with OLS
def do_multivariate_regression_analysis_with_OLS(target_variable_data, independent_variable_data, output_filename):
    # Add a constant term for the intercept
    independent_variable_data['intercept'] = 1

    # Perform multiple linear regression
    X = independent_variable_data
    Y = target_variable_data

    # Perform the regression the ordinal logistic regression model
    model = sm.OLS(Y, X)

    results = model.fit()
    print(results.summary())

    # Get the summary table as a text
    summary_str = results.summary().as_text()

    # Create a new Document
    doc = Document()
    doc.add_heading('Regression Summary', 0)

    # Add the summary to the document
    doc.add_paragraph(summary_str)

    # Save the document
    filename = output_filename
    doc.save(f'{output_path}{filename}')

    return results # .summary()

# interpret_and_recommend_regression function here
def interpret_and_recommend_regression_with_OLS(results, alpha=0.05):
    interpretations = []
    recommendations = []
    
    # Convert the results summary to a DataFrame
    summary_df = pd.read_html(results.summary().tables[1].as_html(), header=0, index_col=0)[0]
    
    # 1. R-squared value
    r_squared = results.rsquared
    msg = f"R^2 (Coefficient of Determination) {r_squared:.3f} "
    msg += f"indicates that approximately {r_squared * 100:.1f}% of the variability in the target variable can be explained by the independent variables."
    interpretations.append(msg)
    
    # Interpret the intercept
    # intercept = summary_df.loc['intercept', 'coef']
    # interpretations.append(f"The estimated intercept is {intercept:.2f}. This is the expected value of the target variable when all independent variables are zero.")

    # 2. Adjusted R^2
    adjusted_r_squared = results.rsquared_adj
    # msg = f" Adjusted R^2 {adjusted_r_squared:.3f} "
    # msg += 'is a more accurate measure when comparing models with different numbers of predictors.'
    # interpretations.append(msg)

    # 3. F-statistic and Prob (F-statistic)
    # Define the significance level
    alpha = 0.05
    f_statistic = results.fvalue
    p_value_f_statistic = results.f_pvalue
    msg = f" F-statistic of ({f_statistic:.3f}) with a Prob (F-statistic) ({p_value_f_statistic:.2e}): "
    if p_value_f_statistic < alpha:
        msg += f"the p-value is less than the significance level of {alpha}, indicating that the model as a whole is statistically significant."
    else:
        msg += f"the p-value is greater than the significance level of {alpha}, indicating that the model as a whole is not statistically significant."        
    interpretations.append(msg)

    # 4. Durbin-Watson statistic
    durbin_watson = sm.stats.stattools.durbin_watson(results.resid)
    # Define the lower and upper bounds for Durbin-Watson statistic
    lower_bound = 1.5
    upper_bound = 2.5
    # Interpret the Durbin-Watson statistic
    if durbin_watson < lower_bound:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests positive autocorrelation."
    elif durbin_watson > upper_bound:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests negative autocorrelation."
    else:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests that there is no autocorrelation."
    interpretations.append(msg)

    # 5. Individual Coefficients and Significant Predictors
    # Define the significance level
    alpha = 0.05

    # Extract coefficients and p-values from the results object
    coefficients = results.params
    p_values = results.pvalues

    # Exclude the intercept (constant term) from the list
    if 'intercept' in coefficients.index:
        coefficients = coefficients.drop('intercept')
        p_values = p_values.drop('intercept')

    # Create a DataFrame with predictors, coefficients, and p-values
    significant_predictors_df = pd.DataFrame({
        'Predictor': coefficients.index,
        'Coefficient': coefficients.values,
        'p-value': p_values.values
    })

    # Filter the DataFrame to include only significant predictors based on the significance level
    significant_predictors_df = significant_predictors_df[significant_predictors_df['p-value'] < alpha]
    # Format the 'Coefficient' and 'p-value' columns to have three decimal places
    significant_predictors_df['Coefficient'] = significant_predictors_df['Coefficient'].apply('{:.3f}'.format)
    significant_predictors_df['p-value'] = significant_predictors_df['p-value'].apply('{:.3f}'.format)

    print(significant_predictors_df)
    # significant_predictors_str = significant_predictors_df.to_string(index=False, header=False)
    significant_predictors_str = '|  '.join(significant_predictors_df.apply(lambda row: ', '.join(row.astype(str)), axis=1))

    msg = "\nIndividual Coefficients:"
    if len(significant_predictors_df) > 0:
        msg += '\nThe following Individual Coefficients with Associated p-values < 0.05 are considered Significant Predictors:\n'
        interpretations.append(msg)
        interpretations.append(significant_predictors_str)

    else:
        msg += '\nNo Individual Coefficients are considered Significant Predictors because their p-values > 0.05.\n'
        interpretations.append(msg)

    print(significant_predictors_str)
    
    interpretation_text = ' '.join(interpretations)
    recommendation_text = ' '.join(recommendations)
    
    return interpretation_text, recommendation_text

# To test a hypothesis: if a factor (with all its constructs/questions) has an effect on the dependent variable.
# Formulate Hypotheses:
# Null Hypothesis (): The coefficients of all the questions representing the "HH" factor are equal to zero (i.e., the "HH" factor has no effect on the dependent variable).
# Alternative Hypothesis (Ha): At least one of the coefficients of the questions representing the "HH" factor is not equal to zero (i.e., the "HH" factor has an effect on the dependent variable).

def test_if_factor_has_effect_on_target(target_variable_data, independent_variable_data, independent_cols):
    # Add a constant term for the intercept
    independent_variable_data['intercept'] = 1
    independent_variable_data_full = independent_variable_data.copy()

    # Perform multiple linear regression
    X = independent_variable_data
    Y = target_variable_data

    # Perform the OLS with FULL independent_variable_data
    model_full = sm.OLS(Y, X)
    results_full = model_full.fit()

    # Fit the reduced model (excluding Factor variables/constructs)
    first_independent_cols_indice = 0
    i = 0
    all_results = []
    for col in independent_cols: 
        independent_variable_data = independent_variable_data_full.copy()
        col_name = col.get('name')
        col_index_from = col.get('col_index_from') 
        col_nbr_questions = col.get('number_questions')
        if i == 0:
            first_independent_cols_indice = col_index_from
        i += 1
        list_cols = []
        for idx in range(col_nbr_questions):
            list_cols.append(col_index_from+idx-first_independent_cols_indice)
        # print(list_cols)
        # Assuming list_cols contains column indices
        column_labels_to_drop = independent_variable_data.columns[list_cols]
        # Drop the specified columns by labels
        X_reduced = independent_variable_data.drop(columns=column_labels_to_drop)
        model_reduced = sm.OLS(Y, X_reduced)
        results_reduced = model_reduced.fit()

        del independent_variable_data

        # Perform the F-test
        f_statistic, p_value, _ = results_full.compare_f_test(results_reduced)
        print(f_statistic, p_value)

        # Check the significance of the F-test
        alpha = 0.05
        if p_value < alpha:
            result = f"The presence of {col_name} has a noticeable impact on the outcome. This means that {col_name} plays a significant role in explaining the variations in the dependent variable (F-statistic: {f_statistic:.3f}, p-value: {p_value:.3f})"        
        else:
            result = f"There is not enough evidence to conclude that the {col_name} factor has a significant effect on the dependent variable (F-statistic: {f_statistic:.3f}, p-value: {p_value:.3f})."
        print(result)

        all_results.append(result)
    
    print(all_results)

    return all_results


# Structural Equation Modeling (SEM) 

# Create latent construct according to variables/cols defined in the data model.
def create_latent_construct(selected_cols):
    constructs =''
    for col in selected_cols:
        construct = format_name(col.get('name'))
        num_items = col.get('number_questions')
        col_index_from = int(col.get('col_index_from'))
        # items = [f"{construct}_Q{col_index_from + i}" for i in range(num_items)]
        items = [f"{construct}_Q{i+1}" for i in range(num_items)]
        constructs += f"{construct} =~ {' + '.join(items)}\n"

    return constructs

def format_name(name):
    """Format the name by removing spaces and special characters."""
    return name.replace(" ", "_").replace("(", "").replace(")", "")

def format_column_names(df):
    """Format the column names of a DataFrame."""
    formatted_columns = {col: format_name(col) for col in df.columns}
    return df.rename(columns=formatted_columns)

def conduct_sem_analysis(independent_cols, target_cols, independent_data, target_data):
    # Format the column names of the independent_data and target_data DataFrames
    print(target_data)
    independent_data = format_column_names(independent_data)
    target_data = format_column_names(target_data)
    new_independent_cols = independent_data.columns
    new_target_cols = target_data.columns
    # print(target_data)
    # print(new_target_cols)
    
    # Generate the measurement model string for independent variables (latent constructs)
    independent_constructs = create_latent_construct(independent_cols)

    # Generate the measurement model string for dependent variables (latent constructs)
    dependent_constructs = create_latent_construct(target_cols)

    # Initialize the structural model string
    structural_model = ""
    
    # Define the relationships between independent and dependent latent constructs
    for target_col in target_cols:
        target_construct = format_name(target_col.get('name'))
        independent_constructs_list = [format_name(col.get('name')) for col in independent_cols]
        structural_model += f"{target_construct} ~ {' + '.join(independent_constructs_list)}\n"
    
    # Combine measurement and structural model strings to create the full SEM syntax string
    sem_model_spec = f'\n{independent_constructs}\n{dependent_constructs}\n{structural_model}\n'

    # Print the generated SEM syntax string
    print(sem_model_spec)

    # Assuming SEM_Model is properly imported and defined
    sem_model = SEM_Model(sem_model_spec)

    # Join the formatted independent_data and target_data DataFrames
    df = independent_data.join(target_data, how='inner')
    # print(df)

    # Fit the SEM model to the survey data
    sem_model.fit(df)

    # Retrieve the results using the inspect method
    results_df = sem_model.inspect()

    # Display the results
    # print(results_df)

    # Define the list of independent factors and dependent variables in your model
    independent_factors = [format_name(col.get('name')) for col in independent_cols]
    dependent_variables = [format_name(col.get('name')) for col in target_cols]

    # Filter results_df to only include rows representing relationships between independent factors and dependent variables
    filtered_results_df = results_df[
        results_df.apply(lambda row: (row['lval'] in independent_factors and row['rval'] in dependent_variables) or
                                      (row['lval'] in dependent_variables and row['rval'] in independent_factors), axis=1)]

    # Display the filtered results
    print(filtered_results_df)

    # Save the Correlation Table to an Excel file
    # Set index=False to exclude the DataFrame index
    filename = 'SEM_Results'
    excel_file_path = f'{output_path}{filename}.xlsx'
    results_df.to_excel(excel_file_path, index=True)  

    return filtered_results_df

# Interpret SEM results
def interpret_sem_results(sem_results):
    # Define significance level
    alpha = 0.05

    # Initialize an empty list to store interpretations
    interpretations = []

    # Iterate over rows of the results DataFrame
    for index, row in sem_results.iterrows():
        # Retrieve relevant information from the row
        lval = row['lval']
        rval = row['rval']
        op = row['op']
        
        # Check if p_value is a valid number
        try:
            p_value = float(row['p-value'])
            estimate = float(row['Estimate'])
        except ValueError:
            print(f"Warning: Unable to convert values to float for {lval} {op} {rval}. Skipping interpretation.")
            continue  # Skip to the next iteration
        
        # Check if the relationship is between a factor and a dependent variable
        # You might need to customize this condition based on your specific model and variable names
        if op == '~':
            relationship = "influences"
        else:
            continue  # Skip rows that do not represent relationships between factors and dependent variables
        
        # Check the significance of the parameter
        if p_value <= alpha:
            significance = "statistically significant"
        else:
            significance = "not statistically significant"

        # Determine the direction of the relationship
        if estimate > 0:
            direction = "positive"
        elif estimate < 0:
            direction = "negative"
        else:
            direction = "neutral"

        # Build the interpretation string
        # interpretation = f"{lval} {relationship} {rval} in a {significance} (p-value= {p_value}) and {direction} (estimate= {estimate}) manner. "
        interpretation = (f"{rval} has a {significance} "
                  f"(p-value= {p_value}) and {direction} "
                  f"(estimate= {estimate}) relationship with {lval}.")
        
        # Append the interpretation to the list
        interpretations.append(interpretation)

    # Print out the interpretations
    #for i, interpretation in enumerate(interpretations, 1):
    #     print(f"Interpretation {i}:", interpretation)

    return interpretations