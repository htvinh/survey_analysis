from docx import Document
import pandas as pd
import statsmodels.api as sm

from sklearn.decomposition import FactorAnalysis
from sklearn.preprocessing import StandardScaler

import numpy as np


import os


output_path = './output/'
# Check if the directory exists
if not os.path.exists(output_path):
    # If it doesn't exist, create the directory
    os.makedirs(output_path)
    print(f"Directory '{output_path}' created.")

def read_excel_to_dict(excel_path):
    # Read the Excel workbook
    xls = pd.ExcelFile(excel_path)
    
    # Initialize an empty dictionary to store all the data
    all_data = {}
    
    # Loop through each sheet and read data into a DataFrame, then convert it to a dictionary
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name)
        all_data[sheet_name] = df.to_dict(orient='index')
        
    return all_data

# Display column names with index
def display_col_name_index(df):
  # Iterate through columns and display column names with indexes in brackets
  for i, column_name in enumerate(df.columns):
      print(f"Column {i} [{column_name}]")

# Rename columns for demographic data columns: to have shorter column names
def rename_demographic_columns_by_index(df, demographic_cols):
    demographic_cols_names = []  # Initialize a list to store independent column names
    for col in demographic_cols:
        col_index = col.get('col_index')
        new_column_name = col.get('name')
        df.rename(columns={df.columns[col_index]: new_column_name}, inplace=True)
        demographic_cols_names.append(new_column_name)
    return df, demographic_cols_names

# Rename Independent columns according to Independent_Cols
def rename_independent_columns(df, independent_cols):
    independent_cols_names = []  # Initialize a list to store independent column names
    for col in independent_cols:
        col_name = col.get('name')
        col_index_from = col.get('col_index_from')
        col_index_to = col_index_from + col.get('number_questions')

        for i in range(col_index_from, col_index_to):
            old_col_name = df.columns[i]
            new_col_name = f'{col_name}_Q{i - col_index_from + 1}'
            df.rename(columns={old_col_name: new_col_name}, inplace=True)
            independent_cols_names.append(new_col_name)  # Append the new column name

    return df, independent_cols_names

# Rename Dependent/Target columns according to Dependent_Cols
def rename_target_columns(df, target_cols):
    target_cols_names = []  # Initialize a list to store Target column names
    for col in target_cols:
        col_name = col.get('name')
        col_index_from = col.get('col_index_from')
        col_index_to = col_index_from + col.get('number_questions')
        # Rename the column
        for i in range(col_index_from, col_index_to):
            old_col_name = df.columns[i]
            new_col_name = f'{col_name}_Q{i - col_index_from + 1}'
            df.rename(columns={old_col_name: new_col_name}, inplace=True)
            target_cols_names.append(new_col_name)  # Append the new column name

    return df, target_cols_names

# Create a statistics table for selected columns
def compute_selected_cols_statistics(df, selected_cols_names, output_file_name):
    nbr_data_points_cleaned = df.shape[0]
    stats_table = {}
    for col_name in selected_cols_names:
        freq = df[col_name].value_counts().sort_index()
        percent = (freq / nbr_data_points_cleaned) * 100
        percent = percent.round(1)  # Round to two decimal places
        stats_table[col_name] = pd.DataFrame({'Frequency': freq, 'Percent': percent})
        
        # Sort the DataFrame by 'Frequency' in descending order
        stats_table[col_name] = stats_table[col_name].sort_values(by='Frequency', ascending=False)
    
    # Save to Excel file
    output_file_name = f'{output_file_name}_Stats'
    save_to_excel(output_file_name, stats_table)

    del df
    return stats_table


# Convert the data in a nested dict into a dictionary of DataFrames
def convert_to_dataframe(table):
    dfs = {}
    for key, value in table.items():
        df = value
        dfs[key] = df

    # Concatenate the DataFrames vertically (along rows)
    combined_df = pd.concat(dfs.values(), axis=1, ignore_index=True)

    return combined_df

# Save the statistics table to an Excel file
def save_to_excel(filename, table):
    with pd.ExcelWriter(f'{output_path}{filename}.xlsx') as writer:
        for col, table in table.items():
            table.to_excel(writer, sheet_name=f'{col}')

# Extract data for selected columns
def extract_selected_colums_data(df, selected_cols_names):
    data_to_extract = df[selected_cols_names]
    data_to_extract = convert_likert_to_numerical(data_to_extract, extended_likert_mapping_all_languages)
    return data_to_extract

# Function to manually convert Likert-scale responses to numerical values
# Function to manually convert Likert-scale responses to numerical values
def convert_likert_to_numerical(df, likert_mapping):
    df_numerical = df.copy()
    # Convert all columns to string type for consistent mapping
    df_numerical = df_numerical.astype(str)
    # Convert the mapping keys to string type as well
    str_likert_mapping = {str(key): value for key, value in likert_mapping.items()}
    
    for column in df.columns:
        df_numerical[column] = df_numerical[column].str.lower().map(str_likert_mapping)
        
    # Check for any remaining non-numeric values
    # if df_numerical.applymap(np.isreal).all().all() == False:
    #     raise AssertionError("Not all columns are numeric after mapping. Check your mapping and DataFrame.")
        
    del df
    return df_numerical




# Sample Likert scale mapping (you can adjust this based on your actual survey Likert scale)
extended_likert_mapping_all_languages = {
    # English Labels
    'strongly disagree': 1,
    'disagree': 2,
    'neutral': 4,
    'agree': 6,
    'strongly agree': 7,
    'no opinion': 4,
   
    # Vietnamese Labels
    'hoàn toàn không đồng ý': 1,
    'không đồng ý': 2,
    'không ý kiến': 4,
    'đồng ý': 6,
    'hoàn toàn đồng ý': 7,
    'trung lap': 4,
    
    # Numerical Levels
    1: 1,
    2: 2,
    3: 3,
    4: 4,
    5: 5,
    6: 6,
    7: 7,
    8: 8,
    9: 9,
}

# Convert dataframe to numerical, keeping ordinality
def convert_df_to_numerical(df):
    # Initialize a list to store the names of detected ordinal columns
    detected_ordinal_columns = []

    # Check each column in the survey data
    for column in df.columns:
        # Get the unique values in the column (excluding NaN)
        unique_values = set(df[column].dropna().unique())
        
        # Check if the unique values are a subset of the extended ordinal mapping keys
        if unique_values.issubset(set(extended_likert_mapping_all_languages.keys())):
            # Column is detected as ordinal, mark it as an ordinal column
            detected_ordinal_columns.append(column)
                
            # Apply the extended ordinal mapping to convert the ordinal responses to numerical values
            df[column] = df[column].map(extended_likert_mapping_all_languages)

    # Identify categorical columns (excluding the detected ordinal columns)
    categorical_columns = df.select_dtypes(include=['object']).columns.difference(detected_ordinal_columns)

    # Convert categorical columns to numerical values using one-hot encoding
    df = pd.get_dummies(df, columns=categorical_columns)

    return df



# Remove data points (rows) with missing data related to selected columns
def remove_rows_with_missing_data_related_to_selected_cols(df, selected_cols_names):
    total_num_data_points = df.shape[0]
    # Filter rows with missing data only in the selected columns
    df = df.dropna(subset=selected_cols_names)
    # Number of data points removed
    number_datapoints_removed = total_num_data_points - df.shape[0]
    print(f'\nNumber of data points removed: \n{number_datapoints_removed}')

    return df

# Testing the reliability of the scale using Cronbach's alpha for independent columns
def cronbach_alpha(items_df):
    """
    Compute Cronbach's Alpha for a set of items (columns) in a DataFrame.
    
    :param items_df: DataFrame, where each column is an item and rows are observations
    :return: float, Cronbach's Alpha
    """
    # Number of items
    item_count = len(items_df.columns)
    
    # Variance for every individual item
    item_variances = items_df.var(axis=0)
    
    # Total variance for the item scores
    total_var = item_variances.sum()
    
    # Sum of the item-pair covariances
    item_covariances_sum = items_df.cov().sum().sum() - total_var
    
    # Cronbach's alpha formula
    alpha = (item_count / (item_count - 1)) * (1 - (total_var / (total_var + item_covariances_sum)))
    
    return alpha

def compute_cronbach_alpha(selected_data, independent_cols):
    # Calculate Cronbach's alpha
    overall_alpha = cronbach_alpha(selected_data)
    overall_alpha = round(overall_alpha, 2)
    overall_alpha_evaluation = interpret_alpha(overall_alpha)

    # Compute Cronbach Alpha for each Independent Variables
    alpha_table = []
    for col in independent_cols:
        col_names = []
        variable_name = col.get('name')
        nbr_questions = col.get('number_questions')
        for idx in range(nbr_questions):
            col_names.append(f'{variable_name}_Q{idx+1}')
        # print(col_names)
        data_of_this_variable = selected_data[col_names]
        # Calcul Alphe for this variable
        alpha_variable = cronbach_alpha(data_of_this_variable)
        alpha_variable = round(alpha_variable, 2)
        # alpha_evaluation = interpret_alpha(alpha_variable[0])
        alpha_evaluation = interpret_alpha(alpha_variable)
        alpha_table.append([alpha_variable, alpha_evaluation])

    # Print the result
    print(f"Overall Cronbach's Alpha: {overall_alpha}")

    return f"{overall_alpha}, '{overall_alpha_evaluation}'", alpha_table

def interpret_alpha(alpha):
    if alpha > 0.9:
        return "Excellent"
    elif alpha > 0.8:
        return "Good"
    elif alpha > 0.7:
        return "Acceptable"
    elif alpha > 0.6:
        return "Questionable"
    elif alpha > 0.5:
        return "Poor"
    else:
        return "Unacceptable"

def conduct_efa_analysis(selected_data):
    # Standardize the data
    scaler = StandardScaler()
    selected_data_standardized = scaler.fit_transform(selected_data)

    # Fit Factor Analysis model to the data to find eigenvalues
    fa = FactorAnalysis()
    fa.fit(selected_data_standardized)
    eigenvalues = np.linalg.eigvals(fa.get_covariance())
    
    # Determine the number of factors to retain based on eigenvalues greater than one
    num_factors = np.sum(eigenvalues > 1)

    # Initialize Factor Analysis
    fa = FactorAnalysis(n_components=num_factors)

    # Fit Factor Analysis to the standardized data
    fa.fit(selected_data_standardized)

    # Loadings are the coefficients of the original variables
    loadings = pd.DataFrame(fa.components_, columns=selected_data.columns)

    # Save the loadings to an Excel file
    filename = 'EFA_Analysis'
    excel_file_path = f'{filename}.xlsx'
    loadings.to_excel(excel_file_path, index=True)  

    return loadings

def interpret_efa_results(efa_results, threshold_high, threshold_moderate):
   # Number of factors extracted
    num_factors = efa_results.shape[0]
    
    # Average loading
    avg_loading = efa_results.abs().mean().mean()

    # Identify questions with low loadings on all factors
    threshold_low = threshold_moderate  # Threshold for low loadings
    low_loading_questions = efa_results.columns[(efa_results.abs() < threshold_low).all(axis=0)]
    
    # Interpretation
    interpretation = []
    interpretation.append(f"The analysis identified {num_factors} latent main factors influencing the responses.")
    
    if avg_loading > threshold_high:
        interpretation.append("Most questions have a strong association with at least one of the identified factors.")
    elif avg_loading > threshold_moderate:
        interpretation.append("There are moderate associations between questions and the identified factors.")
    else:
        interpretation.append("The associations between most questions and the identified factors are relatively weak.")

    if not low_loading_questions.empty:
        interpretation.append(f"The following questions have weak associations with all identified factors and may need to be revised or removed: {', '.join(low_loading_questions)}.")
    
    return interpretation

    
# Pearson correlation matrix
def compute_correlation(selected_data):
    correlation_table = selected_data.corr(method='pearson')
    print("Pearson correlation coefficient matrix: ")
    print(correlation_table)
    # Save the Correlation Table to an Excel file
    # Set index=False to exclude the DataFrame index
    filename = 'Correlation_Table'
    excel_file_path = f'{output_path}{filename}.xlsx'
    correlation_table.to_excel(excel_file_path, index=True)  

    return correlation_table

def interpret_correlation(correlation_table, threshold_strong,threshold_moderate):
    # Remove self-correlations by setting diagonal to NaN
    for i in range(correlation_table.shape[0]):
        correlation_table.iloc[i, i] = np.nan
    
    # Compute the average absolute correlation
    avg_correlation = correlation_table.abs().mean().mean()

    # Identify variables with low correlations with others
    threshold_low = threshold_moderate  # Threshold for low correlation
    low_correlation_variables = correlation_table.columns[(correlation_table.abs().mean() < threshold_low)]
    
    # Identify pairs of variables with high correlations
    threshold_high = threshold_strong  # Threshold for high correlation
    high_correlation_pairs = [(var1, var2) for var1 in correlation_table.columns for var2 in correlation_table.columns if (correlation_table.loc[var1, var2] > threshold_high) and (var1 != var2)]
    
    # Interpretation
    interpretation = []
    if avg_correlation > threshold_strong:
        comment = "Most variables in the dataset are strongly related to each other."
    elif avg_correlation > threshold_moderate:
        comment ="There are moderate relationships between variables in the dataset."
    else:
        comment = "The variables in the dataset are mostly weakly related or unrelated to each other."
    interpretation.append(comment)
    
    if not low_correlation_variables.empty:
        interpretation.append(f"The following variables have weak relationships with most other variables: {', '.join(low_correlation_variables)}. ")
        interpretation.append('Consider reviewing their relevance to the study.')
    if high_correlation_pairs:
        interpretation.append(f"Some pairs of variables are highly correlated, indicating potential redundancy. ") 
        interpretation.append(f"These pairs are: " + ', '.join([f"({var1}, {var2})" for var1, var2 in high_correlation_pairs]) + ".")
        interpretation.append('Consider reviewing whether both variables in each pair are necessary, or whether dimensionality reduction or feature engineering might be appropriate.')
    return interpretation



# Regression Analysis with OLS
def do_multivariate_regression_analysis_with_OLS(target_variable_data, independent_variable_data, output_filename):
    # Add a constant term for the intercept
    independent_variable_data['intercept'] = 1

    # Perform multiple linear regression
    X = independent_variable_data
    Y = target_variable_data

    # Perform the regression the ordinal logistic regression model
    model = sm.OLS(Y, X)

    results = model.fit()
    print(results.summary())

    # Get the summary table as a text
    summary_str = results.summary().as_text()

    # Create a new Document
    doc = Document()
    doc.add_heading('Regression Summary', 0)

    # Add the summary to the document
    doc.add_paragraph(summary_str)

    # Save the document
    filename = output_filename
    doc.save(f'{output_path}{filename}')

    return results # .summary()

# interpret_and_recommend_regression function here
def interpret_and_recommend_regression_with_OLS(results, alpha=0.05):
    interpretations = []
    recommendations = []
    
    # Convert the results summary to a DataFrame
    summary_df = pd.read_html(results.summary().tables[1].as_html(), header=0, index_col=0)[0]
    
    # 1. R-squared value
    r_squared = results.rsquared
    msg = f"R^2 (Coefficient of Determination) {r_squared:.3f} "
    msg += f"indicates that approximately {r_squared * 100:.1f}% of the variability in the target variable can be explained by the independent variables."
    interpretations.append(msg)
    
    # Interpret the intercept
    # intercept = summary_df.loc['intercept', 'coef']
    # interpretations.append(f"The estimated intercept is {intercept:.2f}. This is the expected value of the target variable when all independent variables are zero.")

    # 2. Adjusted R^2
    adjusted_r_squared = results.rsquared_adj
    # msg = f" Adjusted R^2 {adjusted_r_squared:.3f} "
    # msg += 'is a more accurate measure when comparing models with different numbers of predictors.'
    # interpretations.append(msg)

    # 3. F-statistic and Prob (F-statistic)
    # Define the significance level
    alpha = 0.05
    f_statistic = results.fvalue
    p_value_f_statistic = results.f_pvalue
    msg = f" F-statistic of ({f_statistic:.3f}) with a Prob (F-statistic) ({p_value_f_statistic:.2e}): "
    if p_value_f_statistic < alpha:
        msg += f"the p-value is less than the significance level of {alpha}, indicating that the model as a whole is statistically significant."
    else:
        msg += f"the p-value is greater than the significance level of {alpha}, indicating that the model as a whole is not statistically significant."        
    interpretations.append(msg)

    # 4. Durbin-Watson statistic
    durbin_watson = sm.stats.stattools.durbin_watson(results.resid)
    # Define the lower and upper bounds for Durbin-Watson statistic
    lower_bound = 1.5
    upper_bound = 2.5
    # Interpret the Durbin-Watson statistic
    if durbin_watson < lower_bound:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests positive autocorrelation."
    elif durbin_watson > upper_bound:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests negative autocorrelation."
    else:
        msg = f"Durbin-Watson ({durbin_watson:.3f}) suggests that there is no autocorrelation."
    interpretations.append(msg)

    # 5. Individual Coefficients and Significant Predictors
    # Define the significance level
    alpha = 0.05

    # Extract coefficients and p-values from the results object
    coefficients = results.params
    p_values = results.pvalues

    # Exclude the intercept (constant term) from the list
    if 'intercept' in coefficients.index:
        coefficients = coefficients.drop('intercept')
        p_values = p_values.drop('intercept')

    # Create a DataFrame with predictors, coefficients, and p-values
    significant_predictors_df = pd.DataFrame({
        'Predictor': coefficients.index,
        'Coefficient': coefficients.values,
        'p-value': p_values.values
    })

    # Filter the DataFrame to include only significant predictors based on the significance level
    significant_predictors_df = significant_predictors_df[significant_predictors_df['p-value'] < alpha]
    # Format the 'Coefficient' and 'p-value' columns to have three decimal places
    significant_predictors_df['Coefficient'] = significant_predictors_df['Coefficient'].apply('{:.3f}'.format)
    significant_predictors_df['p-value'] = significant_predictors_df['p-value'].apply('{:.3f}'.format)

    print(significant_predictors_df)
    # significant_predictors_str = significant_predictors_df.to_string(index=False, header=False)
    significant_predictors_str = '|  '.join(significant_predictors_df.apply(lambda row: ', '.join(row.astype(str)), axis=1))

    msg = "\nIndividual Coefficients:"
    if len(significant_predictors_df) > 0:
        msg += '\nThe following Individual Coefficients with Associated p-values < 0.05 are considered Significant Predictors:\n'
        interpretations.append(msg)
        interpretations.append(significant_predictors_str)

    else:
        msg += '\nNo Individual Coefficients are considered Significant Predictors because their p-values > 0.05.\n'
        interpretations.append(msg)

    print(significant_predictors_str)
    
    interpretation_text = ' '.join(interpretations)
    recommendation_text = ' '.join(recommendations)
    
    return interpretation_text, recommendation_text

# To test a hypothesis: if a factor (with all its constructs/questions) has an effect on the dependent variable.
# Formulate Hypotheses:
# Null Hypothesis (): The coefficients of all the questions representing the "HH" factor are equal to zero (i.e., the "HH" factor has no effect on the dependent variable).
# Alternative Hypothesis (Ha): At least one of the coefficients of the questions representing the "HH" factor is not equal to zero (i.e., the "HH" factor has an effect on the dependent variable).

def test_if_factor_has_effect_on_target(target_variable_data, independent_variable_data, independent_cols):
    # Add a constant term for the intercept
    independent_variable_data['intercept'] = 1
    independent_variable_data_full = independent_variable_data.copy()

    # Perform multiple linear regression
    X = independent_variable_data
    Y = target_variable_data

    # Perform the OLS with FULL independent_variable_data
    model_full = sm.OLS(Y, X)
    results_full = model_full.fit()

    # Fit the reduced model (excluding Factor variables/constructs)
    first_independent_cols_indice = 0
    i = 0
    all_results = []
    for col in independent_cols: 
        independent_variable_data = independent_variable_data_full.copy()
        col_name = col.get('name')
        col_index_from = col.get('col_index_from') 
        col_nbr_questions = col.get('number_questions')
        if i == 0:
            first_independent_cols_indice = col_index_from
        i += 1
        list_cols = []
        for idx in range(col_nbr_questions):
            list_cols.append(col_index_from+idx-first_independent_cols_indice)
        # print(list_cols)
        # Assuming list_cols contains column indices
        column_labels_to_drop = independent_variable_data.columns[list_cols]
        # Drop the specified columns by labels
        X_reduced = independent_variable_data.drop(columns=column_labels_to_drop)
        model_reduced = sm.OLS(Y, X_reduced)
        results_reduced = model_reduced.fit()

        del independent_variable_data

        # Perform the F-test
        f_statistic, p_value, _ = results_full.compare_f_test(results_reduced)
        print(f_statistic, p_value)

        # Check the significance of the F-test
        alpha = 0.05
        if p_value < alpha:
            result = f"The presence of {col_name} has a noticeable impact on the outcome. This means that {col_name} plays a significant role in explaining the variations in the dependent variable (F-statistic: {f_statistic:.3f}, p-value: {p_value:.3f})"        
        else:
            result = f"There is not enough evidence to conclude that the {col_name} factor has a significant effect on the dependent variable (F-statistic: {f_statistic:.3f}, p-value: {p_value:.3f})."
        print(result)

        all_results.append(result)
    
    print(all_results)

    return all_results

